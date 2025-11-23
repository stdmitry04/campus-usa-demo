# server/documents/ocr_service.py
import logging
import tempfile
import os
from typing import Dict, Optional, Union, Tuple
from django.conf import settings
import requests

import logging
import tempfile
import os
from typing import Dict, Optional, Tuple
from django.conf import settings
import requests

# imports for document parsing
import pdfplumber
from docx import Document as DocxDocument
import mammoth
from PIL import Image
import pytesseract


logger = logging.getLogger(__name__)


class DocumentOCRService:
    """extracts text from various document formats"""

    def __init__(self):
        self.supported_formats = {
            'application/pdf': self._extract_from_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_from_docx,
            'application/msword': self._extract_from_doc,
            'text/plain': self._extract_from_txt,
            'image/jpeg': self._extract_from_image,
            'image/jpg': self._extract_from_image,
            'image/png': self._extract_from_image,
            'image/tiff': self._extract_from_image,
            'image/bmp': self._extract_from_image,
            'application/rtf': self._extract_from_rtf,
        }

    def extract_text_from_document(self, document) -> Dict[str, any]:
        """main method to extract text from a document model instance"""
        try:
            # get file content
            file_content, content_type = self._get_file_content(document)

            if not file_content:
                return {
                    'success': False,
                    'error': 'could not retrieve file content',
                    'text': '',
                    'metadata': {}
                }

            # determine content type
            if not content_type:
                content_type = document.content_type or 'application/octet-stream'

            logger.info(f"extracting text from {content_type} document: {document.title}")

            # find appropriate extraction method
            extractor = self.supported_formats.get(content_type.lower())
            if not extractor:
                return {
                    'success': False,
                    'error': f'unsupported file type: {content_type}',
                    'text': '',
                    'metadata': {'content_type': content_type}
                }

            # extract text using appropriate method
            result = extractor(file_content)

            if result['success']:
                logger.info(f"successfully extracted {len(result['text'])} characters from {document.title}")
            else:
                logger.error(f"failed to extract text from {document.title}: {result['error']}")

            return result

        except Exception as e:
            logger.error(f"unexpected error during text extraction: {e}")
            return {
                'success': False,
                'error': f'extraction failed: {str(e)}',
                'text': '',
                'metadata': {}
            }

    def _get_file_content(self, document) -> Tuple[Optional[bytes], Optional[str]]:
        """get file content from either s3 or local storage"""
        try:
            if document.is_stored_in_s3:
                # get file from s3
                from core.services.storage_service import S3PreSignedURLManager
                s3_manager = S3PreSignedURLManager()

                # get download url
                download_url = s3_manager.generate_download_url(document.s3_key, document.original_filename)
                if not download_url:
                    return None, None

                # download file content
                response = requests.get(download_url, timeout=30)
                response.raise_for_status()

                return response.content, document.content_type

            elif document.file:
                # read from local file
                document.file.seek(0)
                content = document.file.read()
                return content, document.content_type

            else:
                return None, None

        except Exception as e:
            logger.error(f"error retrieving file content: {e}")
            return None, None

    def _extract_from_pdf(self, file_content: bytes) -> Dict[str, any]:
        """extract text from pdf files"""
        if not pdfplumber:
            return {'success': False, 'error': 'pdf processing libraries not available', 'text': '', 'metadata': {}}

        try:
            # create temporary file for pdf processing
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            try:
                # use pdfplumber for better text extraction
                with pdfplumber.open(temp_file_path) as pdf:
                    text_parts = []
                    page_count = len(pdf.pages)

                    for page_num, page in enumerate(pdf.pages, 1):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(f"--- page {page_num} ---\n{page_text}")
                        except Exception as e:
                            logger.warning(f"failed to extract text from page {page_num}: {e}")
                            continue

                    extracted_text = '\n\n'.join(text_parts)

                    return {
                        'success': True,
                        'text': extracted_text,
                        'metadata': {
                            'page_count': page_count,
                            'method': 'pdfplumber',
                            'char_count': len(extracted_text)
                        },
                        'error': None
                    }

            finally:
                # cleanup temp file
                try:
                    os.unlink(temp_file_path)
                except:
                    pass

        except Exception as e:
            logger.error(f"pdf extraction failed: {e}")
            return {'success': False, 'error': str(e), 'text': '', 'metadata': {}}

    def _extract_from_docx(self, file_content: bytes) -> Dict[str, any]:
        """extract text from docx files"""
        if not DocxDocument:
            return {'success': False, 'error': 'docx processing library not available', 'text': '', 'metadata': {}}

        try:
            # create temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            try:
                # open with python-docx
                doc = DocxDocument(temp_file_path)

                # extract paragraphs
                paragraphs = []
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        paragraphs.append(text)

                # extract tables
                table_texts = []
                for table in doc.tables:
                    for row in table.rows:
                        row_texts = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_texts.append(cell_text)
                        if row_texts:
                            table_texts.append(' | '.join(row_texts))

                # combine all text
                all_text_parts = paragraphs
                if table_texts:
                    all_text_parts.append('\n--- tables ---\n')
                    all_text_parts.extend(table_texts)

                extracted_text = '\n'.join(all_text_parts)

                return {
                    'success': True,
                    'text': extracted_text,
                    'metadata': {
                        'paragraph_count': len(paragraphs),
                        'table_count': len(doc.tables),
                        'method': 'python-docx',
                        'char_count': len(extracted_text)
                    },
                    'error': None
                }

            finally:
                try:
                    os.unlink(temp_file_path)
                except:
                    pass

        except Exception as e:
            logger.error(f"docx extraction failed: {e}")
            return {'success': False, 'error': str(e), 'text': '', 'metadata': {}}

    def _extract_from_doc(self, file_content: bytes) -> Dict[str, any]:
        """extract text from legacy doc files using mammoth"""
        if not mammoth:
            return {'success': False, 'error': 'doc processing library not available', 'text': '', 'metadata': {}}

        try:
            # create temporary file
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            try:
                # use mammoth for doc files
                with open(temp_file_path, 'rb') as doc_file:
                    result = mammoth.extract_raw_text(doc_file)

                return {
                    'success': True,
                    'text': result.value,
                    'metadata': {
                        'method': 'mammoth',
                        'char_count': len(result.value),
                        'warnings': [str(warning) for warning in result.messages]
                    },
                    'error': None
                }

            finally:
                try:
                    os.unlink(temp_file_path)
                except:
                    pass

        except Exception as e:
            logger.error(f"doc extraction failed: {e}")
            return {'success': False, 'error': str(e), 'text': '', 'metadata': {}}

    def _extract_from_txt(self, file_content: bytes) -> Dict[str, any]:
        """extract text from plain text files"""
        try:
            # try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    return {
                        'success': True,
                        'text': text,
                        'metadata': {
                            'encoding': encoding,
                            'method': 'direct_decode',
                            'char_count': len(text)
                        },
                        'error': None
                    }
                except UnicodeDecodeError:
                    continue

            return {'success': False, 'error': 'could not decode text file with any encoding', 'text': '', 'metadata': {}}

        except Exception as e:
            logger.error(f"text extraction failed: {e}")
            return {'success': False, 'error': str(e), 'text': '', 'metadata': {}}

    def _extract_from_image(self, file_content: bytes) -> Dict[str, any]:
        """extract text from images using ocr"""
        if not Image or not pytesseract:
            return {'success': False, 'error': 'image processing libraries not available', 'text': '', 'metadata': {}}

        try:
            # create temporary file
            with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            try:
                # open image with pillow
                image = Image.open(temp_file_path)

                # convert to rgb if needed
                if image.mode != 'RGB':
                    image = image.convert('RGB')

                # extract text using tesseract
                extracted_text = pytesseract.image_to_string(image, lang='eng')

                return {
                    'success': True,
                    'text': extracted_text,
                    'metadata': {
                        'image_size': image.size,
                        'image_mode': image.mode,
                        'method': 'tesseract_ocr',
                        'char_count': len(extracted_text)
                    },
                    'error': None
                }

            finally:
                try:
                    os.unlink(temp_file_path)
                except:
                    pass

        except Exception as e:
            logger.error(f"image ocr failed: {e}")
            return {'success': False, 'error': str(e), 'text': '', 'metadata': {}}

    def _extract_from_rtf(self, file_content: bytes) -> Dict[str, any]:
        """extract text from rtf files - basic implementation"""
        try:
            # simple rtf text extraction - removes most formatting
            text = file_content.decode('utf-8', errors='ignore')

            # remove rtf control sequences (very basic)
            import re
            # remove rtf header and control words
            text = re.sub(r'\\[a-z]+\d*', '', text)
            text = re.sub(r'[{}]', '', text)
            text = re.sub(r'\\', '', text)

            # clean up whitespace
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            cleaned_text = '\n'.join(lines)

            return {
                'success': True,
                'text': cleaned_text,
                'metadata': {
                    'method': 'basic_rtf_parser',
                    'char_count': len(cleaned_text)
                },
                'error': None
            }

        except Exception as e:
            logger.error(f"rtf extraction failed: {e}")
            return {'success': False, 'error': str(e), 'text': '', 'metadata': {}}

    def get_supported_formats(self) -> list:
        """return list of supported content types"""
        available_formats = []

        for content_type in self.supported_formats.keys():
            # check if required libraries are available
            if content_type.startswith('application/pdf') and pdfplumber:
                available_formats.append(content_type)
            elif content_type.endswith('wordprocessingml.document') and DocxDocument:
                available_formats.append(content_type)
            elif content_type == 'application/msword' and mammoth:
                available_formats.append(content_type)
            elif content_type == 'text/plain':
                available_formats.append(content_type)
            elif content_type.startswith('image/') and Image and pytesseract:
                available_formats.append(content_type)
            elif content_type == 'application/rtf':
                available_formats.append(content_type)

        return available_formats